import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Título
title = doc.add_heading('Informe Técnico del Proyecto: Sistema de Microservicios', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introducción
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    "El presente informe detalla la arquitectura, estructura y funcionamiento del sistema basado en microservicios del proyecto. "
    "El sistema está diseñado para ser escalable, modular y mantenible, utilizando Spring Boot para el backend, Docker para la "
    "contenedorización y PostgreSQL como sistema de gestión de bases de datos."
)

# Arquitectura
doc.add_heading('2. Arquitectura de Microservicios', level=1)
doc.add_paragraph("El proyecto se compone de los siguientes microservicios principales:")

services = {
    "Eureka Server": "Actúa como servidor de registro y descubrimiento. Permite que todos los demás microservicios se registren bajo un nombre y puedan encontrarse entre sí de manera dinámica, sin necesidad de conocer las direcciones IP o puertos específicos de antemano.",
    "ms-usuarios": "Gestiona la información de los usuarios del sistema, incluyendo la autenticación, autorización y administración de perfiles.",
    "ms-productos": "Encargado de la gestión del catálogo de productos, sus detalles, categorías y precios.",
    "ms-inventario": "Controla el stock disponible de los productos en tiempo real.",
    "ms-tiendas": "Administra la información de las sucursales o tiendas físicas asociadas al negocio.",
    "ms-carrito": "Maneja la lógica del carrito de compras de los usuarios, almacenando temporalmente los productos seleccionados antes del pago.",
    "ms-pedidos": "Procesa las órdenes de compra consolidadas por el usuario, coordinando la confirmación de la venta.",
    "ms-logistica": "Gestiona el envío y seguimiento de los paquetes hacia el destino del cliente final.",
    "ms-promociones": "Aplica descuentos y reglas de negocio promocionales a los productos o al carrito de compras.",
    "ms-facturacion": "Genera comprobantes de pago (facturas o boletas) para las transacciones completadas.",
    "ms-soporte": "Maneja el centro de ayuda, atención de reclamos y tickets de soporte técnico de los clientes."
}

for name, desc in services.items():
    p = doc.add_paragraph()
    p.add_run(f"- {name}: ").bold = True
    p.add_run(desc)

# Estructura de carpetas
doc.add_heading('3. Estructura Interna de Carpetas (Patrón de Diseño)', level=1)
doc.add_paragraph("Cada microservicio sigue una arquitectura multicapa estándar de Spring Boot, organizada en los siguientes paquetes:")

folders = {
    "Controller": "Es el punto de entrada de las peticiones HTTP (API REST). Recibe las solicitudes del cliente (como Postman o el Frontend), valida los parámetros de entrada y delega la lógica al Service. Define las rutas y los métodos (GET, POST, PUT, DELETE).",
    "Service": "Contiene toda la lógica de negocio del microservicio. Se asegura de que se cumplan las reglas de la aplicación antes de persistir la información. Es el intermediario entre el Controller y el Repository.",
    "Repository": "Es la capa de acceso a datos. Utiliza Spring Data JPA para comunicarse con la base de datos (PostgreSQL). Define los métodos para guardar, actualizar, eliminar y consultar registros (consultas SQL abstraídas).",
    "Model (Entities)": "Contiene las clases que representan las tablas de la base de datos. Se mapean usando anotaciones JPA (como @Entity, @Table, @Column) para definir la estructura de la información almacenada.",
    "Exceptions": "Centraliza el manejo de errores. Contiene excepciones personalizadas (ej. RecursoNoEncontradoException) y manejadores globales (@ControllerAdvice) para asegurar que el cliente reciba respuestas HTTP consistentes cuando ocurre un error.",
    "Security": "Almacena las clases de configuración de seguridad, filtros, encriptación de contraseñas y la gestión de acceso basada en roles o permisos de los usuarios."
}

for name, desc in folders.items():
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)

# DTOs
doc.add_heading('4. Uso de DTOs (Data Transfer Objects)', level=1)
doc.add_paragraph(
    "Los DTOs son objetos que se utilizan para encapsular datos y enviarlos de una capa a otra (o hacia el cliente externo). "
    "Su principal función es evitar la exposición de las Entidades de la base de datos (Models) directamente a través del Controller. "
    "Esto mejora la seguridad, oculta información sensible (como contraseñas) y permite enviar solo los datos estrictamente "
    "necesarios para una vista en particular."
)

# JWT
doc.add_heading('5. Seguridad y Autenticación con JWT (JSON Web Token)', level=1)
doc.add_paragraph(
    "El proyecto implementa seguridad a través de JWT. Cuando un usuario se autentica correctamente en el ms-usuarios, "
    "se genera un token JWT firmado mediante una clave secreta configurada en el archivo de propiedades. Este token "
    "contiene información cifrada (claims) sobre la identidad y roles del usuario. "
    "En peticiones posteriores, el cliente debe enviar este token en el header 'Authorization', permitiendo a los microservicios "
    "verificar su identidad y conceder acceso a los recursos protegidos."
)

# Base de datos
doc.add_heading('6. Conexión a la Base de Datos', level=1)
doc.add_paragraph(
    "Cada microservicio se conecta a su propia base de datos (PostgreSQL) garantizando la independencia e integridad de los datos. "
    "La configuración se define en el archivo application.yml (o properties), especificando el driver, la URL de conexión, "
    "el usuario y contraseña, y el comportamiento de Hibernate (por ejemplo, ddl-auto: update). "
    "Todos los esquemas de bases de datos son creados a partir de scripts SQL alojados en el directorio 'init-multi-db'."
)

# Docker
doc.add_heading('7. Despliegue con Docker', level=1)
doc.add_paragraph(
    "El proyecto hace uso de Docker para crear entornos consistentes e independientes. "
    "La carpeta 'init-multi-db' contiene los archivos y scripts (como archivos .bat y .sql) necesarios para levantar e "
    "inicializar un contenedor principal de PostgreSQL con múltiples bases de datos ya preparadas para cada microservicio. "
    "Esto permite que con un solo comando se pueda montar todo el ecosistema de base de datos necesario para el funcionamiento del proyecto."
)

# Postman
doc.add_heading('8. Comunicación y Pruebas con Postman', level=1)
doc.add_paragraph(
    "Postman se utiliza como herramienta principal para simular y validar la comunicación de los clientes con el backend. "
    "A través de Postman, se envían peticiones HTTP hacia los diferentes endpoints de cada microservicio o a través de un API Gateway. "
    "El proyecto incluye colecciones de Postman (archivos .json como Pruebas_Seguridad_Postman.json) preparadas con variables de entorno y scripts "
    "que permiten automatizar las pruebas, como la obtención de un JWT al iniciar sesión y su posterior inyección automática en las cabeceras "
    "para probar endpoints protegidos."
)

doc.save('Informe_Proyecto_Microservicios.docx')
print("Documento Word generado exitosamente: Informe_Proyecto_Microservicios.docx")
